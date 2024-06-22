import whisper
from docx import Document
import torch


def save_transcription_as_docx(transcription_text, filename):
    doc = Document()

    # Add transcription text as a heading and paragraph
    doc.add_heading('Transcription', level=1)
    doc.add_paragraph(transcription_text)

    doc.save(filename)


# Load the largest Whisper model for better accuracy
model = whisper.load_model("large", device="cpu")  # Force the model to run on the CPU

# Load audio and pad/trim it to fit 30 seconds
audio = whisper.load_audio("data/bilingual.mp3")
audio = whisper.pad_or_trim(audio)

# Make log-Mel spectrogram and move to the same device as the model
mel = whisper.log_mel_spectrogram(audio).to("cpu")  # Ensure the mel spectrogram is on the CPU

# Debug: Print the shape of the mel spectrogram
print(f"Mel spectrogram shape: {mel.shape}")

# Decode the audio with options, allowing the model to handle multiple languages
options = whisper.DecodingOptions(fp16=False, language=None)
result = whisper.decode(model, mel, options)

# Print the recognized text
print(result.text)

save_transcription_as_docx(result.text, 'data/bilingual.docx')
