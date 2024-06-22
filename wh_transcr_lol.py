import whisper
from docx import Document
import os

def save_as_docx(minutes, filename):
    doc = Document()
    for key, value in minutes.items():
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        doc.add_paragraph()
    doc.save(filename)


def save_transcription_as_docx(transcription_text, filename):
    doc = Document()

    # Add transcription text as a heading and paragraph
    doc.add_heading('Transcription', level=1)
    doc.add_paragraph(transcription_text)

    doc.save(filename)


# Load a larger Whisper model
model = whisper.load_model("medium")

# Load audio and pad/trim it to fit 30 seconds
audio = whisper.load_audio("data/bilingual.mp4")
audio = whisper.pad_or_trim(audio)

# Make log-Mel spectrogram and move to the same device as the model
mel = whisper.log_mel_spectrogram(audio).to(model.device)

# Detect the spoken language
_, probs = model.detect_language(mel)
detected_language = max(probs, key=probs.get)
print(f"Detected language: {detected_language}")

# Decode the audio with options
options = whisper.DecodingOptions(fp16=False)  # Set fp16 to False if running on CPU
result = whisper.decode(model, mel, options)

# Print the recognized text
print(result.text)

save_transcription_as_docx(result.text, 'data/bilingual.docx')